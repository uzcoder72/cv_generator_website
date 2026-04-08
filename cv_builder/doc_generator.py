from docx import Document
from django.http import HttpResponse
from io import BytesIO

def create_doc_response(resume, template_id):
    doc = Document()
    
    # Build the Word Document via Python!
    doc.add_heading(f"{resume.name.upper()} {resume.surname.upper()}", 0)
    doc.add_paragraph(f"{resume.specialization} | {resume.email} | {resume.contacts}")

    doc.add_heading('Experience', level=1)
    doc.add_paragraph(f"• {resume.exp_position} at {resume.exp_employer}")
    doc.add_paragraph(f"  {resume.exp_start} to {resume.exp_end}")
    
    doc.add_heading('Education', level=1)
    doc.add_paragraph(f"• {resume.edu_degree} from {resume.edu_institution}")
    doc.add_paragraph(f"  {resume.edu_start} to {resume.edu_end}")
    
    doc.add_heading('Skills & Languages', level=1)
    doc.add_paragraph(f"Skills: {resume.skills}")
    doc.add_paragraph(f"Languages: {resume.languages}")
    
    if resume.extra_info:
        doc.add_heading('Extra Information', level=1)
        doc.add_paragraph(resume.extra_info)
    
    # Save to PC memory so Django can send it over HTTP
    f = BytesIO()
    doc.save(f)
    f.seek(0)
    
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{resume.name}_{resume.surname}_CV.docx"'
    return response
